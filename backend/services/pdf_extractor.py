"""
PDF Text Extractor — Reusable Utility
Emirati Journey Platform — Qwen Migration

Standalone PDF/DOCX/TXT extraction using pdfplumber.
Designed to be imported by cv_parser.py and resume_parser.py alike.
Handles bilingual Arabic/English, multi-column layouts, and tables.
"""

import base64
import logging
import os
import re
from typing import Optional

logger = logging.getLogger(__name__)

# Lazy imports to avoid hard failures if libraries are missing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logger.warning("pdfplumber not installed — PDF extraction will be unavailable")

try:
    import fitz as pymupdf  # PyMuPDF — fallback for image-heavy PDFs
except ImportError:
    pymupdf = None
    logger.warning("PyMuPDF not installed — fallback PDF extraction unavailable")

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    logger.warning("python-docx not installed — DOCX extraction will be unavailable")


# Image formats we accept for OCR, and the MIME type each needs in the data
# URL. Derived as one mapping so the accepted list and the MIME lookup cannot
# drift apart — an extension present in one but not the other would either be
# rejected at the door or sent with the wrong MIME.
IMAGE_MIME_TYPES = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
}
IMAGE_EXTENSIONS = tuple(IMAGE_MIME_TYPES)

# Pages OCR'd per PDF. This is a COST cap against the paid DashScope endpoint,
# not a technical limit — named here so it is one edit to raise if OCR ever
# moves onto our own GPU.
OCR_MAX_PAGES = 5

# Longest edge, in pixels, an image is scaled to before upload. Phone photos
# arrive at 8-12 MP; the payload is slow to send while the model downsamples to
# max_pixels regardless.
MAX_IMAGE_EDGE = 2000

# The Arabic instruction is deliberate: these are bilingual documents, and
# without it the model transliterates or drops Arabic blocks.
OCR_PROMPT = (
    "Read all the text in this image. Output the raw text only, "
    "preserving layout. Include Arabic text as-is."
)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    """Extract clean text from a PDF, DOCX, TXT, or image file.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        Extracted text preserving reading order, or empty string on failure.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext == ".txt":
        return _extract_txt(file_path)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_image(file_path)
    else:
        logger.warning(f"Unsupported file extension: {ext}")
        return ""


def extract_text_from_stream(file_stream, filename: str = "") -> str:
    """Extract text from an in-memory file stream.

    Args:
        file_stream: A file-like object (e.g. from Flask request.files).
        filename: Original filename used to detect type.

    Returns:
        Extracted text or empty string.
    """
    ext = os.path.splitext(filename)[1].lower() if filename else ".pdf"

    if ext == ".pdf":
        return _extract_pdf_stream(file_stream)
    elif ext in (".docx", ".doc"):
        return _extract_docx_stream(file_stream)
    elif ext == ".txt":
        return file_stream.read().decode("utf-8", errors="ignore")
    elif ext in IMAGE_EXTENSIONS:
        file_stream.seek(0)
        return _extract_image_bytes(file_stream.read(), filename or ext)
    else:
        logger.warning(f"Unsupported stream file extension: {ext}")
        return ""


# ---------------------------------------------------------------------------
# PDF Extraction (pdfplumber)
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file on disk.

    Strategy:
    1. Try pdfplumber (best for selectable-text PDFs)
    2. If empty, try PyMuPDF/fitz (handles more embedded fonts)
    3. If still empty, use Qwen Vision OCR (scanned/image-based PDFs)
    """
    text = ""

    # Strategy 1: pdfplumber
    if pdfplumber:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = _process_pdf_pages(pdf)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed ({file_path}): {e}")

    # Strategy 2: PyMuPDF fallback (handles embedded fonts)
    if len(text.strip()) < 50 and pymupdf:
        logger.info(f"pdfplumber returned {len(text)} chars — trying PyMuPDF fallback")
        try:
            text = _extract_pdf_pymupdf(file_path)
        except Exception as e:
            logger.warning(f"PyMuPDF extraction also failed ({file_path}): {e}")

    # Strategy 3: Vision OCR via Qwen (scanned/image-based PDFs)
    if len(text.strip()) < 50 and pymupdf:
        logger.info(f"Text extraction returned {len(text)} chars — trying Vision OCR")
        try:
            text = _extract_pdf_vision_ocr(file_path)
        except Exception as e:
            logger.warning(f"Vision OCR extraction failed ({file_path}): {e}")

    if not text.strip():
        logger.error(f"All PDF extraction strategies returned empty for {file_path}")

    return text


def _extract_pdf_stream(file_stream) -> str:
    """Extract text from a PDF file stream.

    Strategy: identical to _extract_pdf.
    1. Try pdfplumber (best for selectable-text PDFs)
    2. If empty, try PyMuPDF/fitz (handles more embedded fonts)
    3. If still empty, use Qwen Vision OCR (scanned/image-based PDFs)

    The two functions must stay in step. They had drifted: this one stopped
    after strategy 2 and returned "" for any scanned PDF, so whether OCR ran
    depended on which entry point a caller happened to reach for.
    """
    text = ""

    # Strategy 1: pdfplumber
    if pdfplumber:
        try:
            file_stream.seek(0)
            with pdfplumber.open(file_stream) as pdf:
                text = _process_pdf_pages(pdf)
        except Exception as e:
            logger.warning(f"pdfplumber stream extraction failed: {e}")

    # Strategy 2: PyMuPDF fallback
    if len(text.strip()) < 50 and pymupdf:
        logger.info(f"pdfplumber stream returned {len(text)} chars — trying PyMuPDF")
        try:
            file_stream.seek(0)
            data = file_stream.read()
            doc = pymupdf.open(stream=data, filetype="pdf")
            parts = []
            for page in doc:
                page_text = page.get_text("text")
                if page_text:
                    parts.append(page_text)
            doc.close()
            text = _clean_text("\n".join(parts))
        except Exception as e:
            logger.warning(f"PyMuPDF stream extraction also failed: {e}")

    # Strategy 3: Vision OCR via Qwen (scanned/image-based PDFs)
    if len(text.strip()) < 50 and pymupdf:
        logger.info(f"Stream text extraction returned {len(text)} chars — trying Vision OCR")
        try:
            file_stream.seek(0)
            text = _ocr_pdf_bytes(file_stream.read())
        except Exception as e:
            logger.warning(f"Vision OCR stream extraction failed: {e}")

    if not text.strip():
        logger.error("All PDF stream extraction strategies returned empty")

    return text


def _extract_pdf_pymupdf(file_path: str) -> str:
    """Fallback PDF extraction using PyMuPDF (fitz).

    PyMuPDF handles embedded fonts, Type3 fonts, and some image-based
    content better than pdfplumber. It's also faster for large PDFs.
    """
    doc = pymupdf.open(file_path)
    parts: list[str] = []

    for page in doc:
        # get_text("text") extracts in reading order
        page_text = page.get_text("text")
        if page_text and page_text.strip():
            parts.append(page_text)

    page_count = len(doc)
    doc.close()
    text = "\n".join(parts).strip()
    text = _clean_text(text)
    logger.info(f"PyMuPDF extraction: {len(text)} chars from {page_count} pages")
    return text


def _ocr_client():
    """The vision client used for OCR, or None if OCR is unavailable.

    Same endpoint rules as qwen_client: the on-premises balancer needs no
    vendor key and must bypass the corporate proxy; DashScope needs both.
    Returns None rather than raising: OCR is a fallback, so a missing key or
    package should degrade to "no text" and let the caller carry on, not break
    an upload.
    """
    try:
        from backend.config.qwen_config import QWEN_API_KEY, QWEN_BASE_URL, QWEN_IS_LOCAL
    except ImportError:  # pragma: no cover — the app runs under both roots
        from config.qwen_config import QWEN_API_KEY, QWEN_BASE_URL, QWEN_IS_LOCAL
    if not QWEN_API_KEY:
        logger.warning("No API key and no local endpoint — Vision OCR unavailable")
        return None
    try:
        from openai import OpenAI
    except ImportError:
        logger.warning("openai package not installed — Vision OCR unavailable")
        return None
    kwargs = {"api_key": QWEN_API_KEY, "base_url": QWEN_BASE_URL}
    proxy = os.getenv("HTTPS_PROXY") or os.getenv("https_proxy") or os.getenv("HTTP_PROXY") or os.getenv("http_proxy")
    if proxy and not QWEN_IS_LOCAL:
        import httpx
        kwargs["http_client"] = httpx.Client(proxy=proxy, timeout=120)
    return OpenAI(**kwargs)


def ocr_image_content(b64_img: str, mime: str) -> list:
    """The multimodal message content for one image. DashScope's qwen-vl-ocr
    accepts min/max pixel hints inside the image part; the OpenAI schema vLLM
    enforces does not, so they are sent only to DashScope."""
    try:
        from backend.config.qwen_config import QWEN_IS_LOCAL
    except ImportError:  # pragma: no cover
        from config.qwen_config import QWEN_IS_LOCAL
    image = {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64_img}"}}
    if not QWEN_IS_LOCAL:
        image.update({"min_pixels": 28 * 28 * 4, "max_pixels": 1280 * 784})
    return [image, {"type": "text", "text": OCR_PROMPT}]


def _ocr_image_bytes(img_bytes: bytes, client, label: str = "image",
                     mime: str = "image/png") -> str:
    """OCR one image through the vision model (qwen-vl-ocr on DashScope, the
    local Qwen3.8 on the balancer — it has native vision).

    Args:
        img_bytes: Raw image bytes.
        client: Client from _ocr_client().
        label: Identifies the page or file in logs.
        mime: MIME type matching img_bytes — a wrong one is rejected upstream.

    Returns:
        Extracted text, or "" if the call failed.
    """
    b64_img = base64.b64encode(img_bytes).decode("utf-8")
    try:
        from backend.config.qwen_config import QWEN_VISION_MODEL
    except ImportError:  # pragma: no cover
        from config.qwen_config import QWEN_VISION_MODEL
    try:
        response = client.chat.completions.create(
            model=QWEN_VISION_MODEL,
            messages=[{"role": "user", "content": ocr_image_content(b64_img, mime)}],
            max_tokens=4096,
            temperature=0,
        )
        page_text = response.choices[0].message.content
        if page_text and page_text.strip():
            logger.info(f"Vision OCR {label}: {len(page_text)} chars extracted")
            return page_text.strip()
    except Exception as ocr_err:
        logger.warning(f"Vision OCR {label} failed: {ocr_err}")

    return ""


def _ocr_pdf_document(doc, max_pages: int) -> str:
    """OCR the first max_pages pages of an open PyMuPDF document."""
    client = _ocr_client()
    if client is None:
        return ""

    page_count = min(len(doc), max_pages)
    all_text_parts: list[str] = []

    for i in range(page_count):
        # Render page at 200 DPI for good OCR quality
        mat = pymupdf.Matrix(200 / 72, 200 / 72)
        pix = doc[i].get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")

        logger.info(f"Vision OCR: processing page {i + 1}/{page_count} ({len(img_bytes)} bytes)")
        page_text = _ocr_image_bytes(img_bytes, client, label=f"page {i + 1}")
        if page_text:
            all_text_parts.append(page_text)

    text = _clean_text("\n\n".join(all_text_parts))
    logger.info(f"Vision OCR total: {len(text)} chars from {page_count} pages")
    return text


def _extract_pdf_vision_ocr(file_path: str, max_pages: int = OCR_MAX_PAGES) -> str:
    """Extract text from a scanned/image-based PDF on disk using Qwen Vision OCR."""
    doc = pymupdf.open(file_path)
    try:
        return _ocr_pdf_document(doc, max_pages)
    finally:
        doc.close()


def _ocr_pdf_bytes(data: bytes, max_pages: int = OCR_MAX_PAGES) -> str:
    """Extract text from a scanned/image-based PDF held in memory.

    The stream counterpart of _extract_pdf_vision_ocr. Both delegate to
    _ocr_pdf_document so the two entry points cannot drift again.
    """
    doc = pymupdf.open(stream=data, filetype="pdf")
    try:
        return _ocr_pdf_document(doc, max_pages)
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# Images (photographed and scanned documents)
# ---------------------------------------------------------------------------

def _prepare_image(data: bytes, ext: str):
    """Normalise an image to PNG, downscaling oversized photos.

    Returns (bytes, mime). Falls back to the original bytes with the MIME type
    matching its extension if PyMuPDF cannot decode the format — sending the
    original is better than failing the extraction outright.
    """
    if pymupdf and ext:
        try:
            doc = pymupdf.open(stream=data, filetype=ext.lstrip("."))
            try:
                page = doc[0]
                longest = max(page.rect.width, page.rect.height)
                scale = min(1.0, MAX_IMAGE_EDGE / longest) if longest else 1.0
                pix = page.get_pixmap(matrix=pymupdf.Matrix(scale, scale))
                return pix.tobytes("png"), "image/png"
            finally:
                doc.close()
        except Exception as e:
            logger.debug(f"Image normalisation failed ({ext}) — sending original: {e}")

    return data, IMAGE_MIME_TYPES.get(ext, "image/png")


def _extract_image_bytes(data: bytes, label: str = "image") -> str:
    """OCR a photographed or scanned document supplied as image bytes.

    Args:
        data: Raw image bytes.
        label: Filename (or any identifier); its extension selects the MIME type.

    Returns:
        Extracted text, or "" if OCR is unavailable or found nothing.
    """
    client = _ocr_client()
    if client is None:
        return ""

    ext = os.path.splitext(label)[1].lower()
    payload, mime = _prepare_image(data, ext)
    text = _clean_text(_ocr_image_bytes(payload, client, label=label, mime=mime))

    if not text.strip():
        logger.error(f"Image extraction returned empty for {label}")
    return text


def _extract_image(file_path: str) -> str:
    """OCR an image file on disk.

    A photographed certificate is the commonest scan we receive, and before
    this it was refused at the extension check — no OCR was ever attempted.
    """
    try:
        with open(file_path, "rb") as fh:
            data = fh.read()
    except OSError as e:
        logger.warning(f"Could not read image {file_path}: {e}")
        return ""

    return _extract_image_bytes(data, os.path.basename(file_path))


def _process_pdf_pages(pdf) -> str:
    """Process all pages of an opened pdfplumber PDF object.

    Extracts body text AND tabular data for comprehensive coverage.
    """
    parts: list[str] = []

    for i, page in enumerate(pdf.pages):
        try:
            # 1. Body text (preserves multi-column reading order)
            body = page.extract_text(x_tolerance=3, y_tolerance=3)
            if body:
                parts.append(body)

            # 2. Tables — convert rows to pipe-delimited strings
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        row_text = " | ".join(cell or "" for cell in row).strip()
                        if row_text and row_text != "|":
                            parts.append(row_text)

        except Exception as page_err:
            logger.warning(f"Error extracting page {i + 1}: {page_err}")

    text = "\n".join(parts).strip()
    text = _clean_text(text)
    logger.info(f"PDF extraction complete: {len(text)} chars from {len(pdf.pages)} pages")
    return text


# ---------------------------------------------------------------------------
# DOCX Extraction
# ---------------------------------------------------------------------------

def _extract_docx(file_path: str) -> str:
    if not DocxDocument:
        logger.error("python-docx not available for DOCX extraction")
        return ""
    try:
        doc = DocxDocument(file_path)
        return _process_docx(doc)
    except Exception as e:
        logger.error(f"DOCX extraction error ({file_path}): {e}")
        return ""


def _extract_docx_stream(file_stream) -> str:
    if not DocxDocument:
        logger.error("python-docx not available for DOCX extraction")
        return ""
    try:
        file_stream.seek(0)
        doc = DocxDocument(file_stream)
        return _process_docx(doc)
    except Exception as e:
        logger.error(f"DOCX stream extraction error: {e}")
        return ""


def _process_docx(doc) -> str:
    """Extract paragraphs and table cells from a python-docx Document."""
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    return _clean_text(text)


# ---------------------------------------------------------------------------
# TXT Extraction
# ---------------------------------------------------------------------------

def _extract_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return _clean_text(f.read())
    except Exception as e:
        logger.error(f"TXT extraction error ({file_path}): {e}")
        return ""


# ---------------------------------------------------------------------------
# Text Cleaning
# ---------------------------------------------------------------------------

# Control characters that should never appear in clean text
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean_text(text: str) -> str:
    """Remove control characters and normalise whitespace."""
    text = _CONTROL_CHARS.sub("", text)
    # Collapse 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
