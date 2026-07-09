"""
CV Export Module
Handles exporting CVs to various formats (PDF, DOCX, JSON)
"""

import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any
import os
from pathlib import Path
import tempfile

# Import libraries for document generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. PDF export will be limited.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX export will be limited.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def get_field(obj: Dict, *keys, default=None):
    """
    Get a field from a dict, trying multiple key names (camelCase and snake_case).
    Returns the first found value or the default.
    """
    if not isinstance(obj, dict):
        return default
    for key in keys:
        if key in obj and obj[key] is not None:
            return obj[key]
    return default


class CVExporter:
    """
    CV Export engine supporting multiple formats
    """
    
    def __init__(self):
        """Initialize the CV Exporter"""
        self.exports_dir = Path(__file__).parent.parent / "exports"
        self.exports_dir.mkdir(exist_ok=True)
        
        # UAE-specific styling - using teal/green theme to match preview
        self.uae_colors = {
            'primary': '#0D9488',    # Teal (matches preview)
            'secondary': '#0D9488',  # Teal
            'accent': '#FFD700',     # Gold
            'text': '#2C3E50',       # Dark blue-gray
            'light_gray': '#F8F9FA'
        }
        
        logger.info("CV Exporter initialized")
    
    def export_cv(self, cv_data: Dict[str, Any], format: str) -> Optional[str]:
        """
        Export CV to specified format
        
        Args:
            cv_data: Complete CV data structure
            format: Export format ('pdf', 'docx', 'json')
            
        Returns:
            Path to exported file or None if failed
        """
        try:
            cv_id = cv_data['metadata']['cv_id']
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            if format == 'json':
                return self._export_json(cv_data, cv_id, timestamp)
            elif format == 'pdf':
                return self._export_pdf(cv_data, cv_id, timestamp)
            elif format == 'docx':
                return self._export_docx(cv_data, cv_id, timestamp)
            else:
                raise ValueError(f"Unsupported export format: {format}")
                
        except Exception as e:
            logger.error(f"Error exporting CV to {format}: {str(e)}")
            raise e
    
    def _export_json(self, cv_data: Dict[str, Any], cv_id: str, timestamp: str) -> str:
        """Export CV as JSON file"""
        try:
            filename = f"cv_{cv_id}_{timestamp}.json"
            file_path = self.exports_dir / filename
            
            # Clean and format the data
            export_data = {
                'metadata': cv_data['metadata'],
                'data': cv_data['data'],
                'exported_at': datetime.now().isoformat(),
                'export_format': 'json',
                'version': '1.0'
            }
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Exported CV {cv_id} to JSON: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error exporting JSON: {str(e)}")
            raise
    
    def _export_pdf(self, cv_data: Dict[str, Any], cv_id: str, timestamp: str) -> str:
        """Export CV as PDF file"""
        if not REPORTLAB_AVAILABLE:
            return self._export_simple_pdf(cv_data, cv_id, timestamp)
        
        try:
            filename = f"cv_{cv_id}_{timestamp}.pdf"
            file_path = self.exports_dir / filename
            
            # Create PDF document
            doc = SimpleDocTemplate(
                str(file_path),
                pagesize=A4,
                rightMargin=50,
                leftMargin=50,
                topMargin=50,
                bottomMargin=30
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles matching the preview
            name_style = ParagraphStyle(
                'NameStyle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=6,
                textColor=colors.HexColor('#1F2937'),
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=20,
                textColor=colors.HexColor('#6B7280'),
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'SectionHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=16,
                spaceAfter=10,
                textColor=colors.HexColor(self.uae_colors['primary']),
                borderWidth=0,
                borderPadding=0,
                leftIndent=0,
                fontName='Helvetica-Bold'
            )
            
            job_title_style = ParagraphStyle(
                'JobTitle',
                parent=styles['Normal'],
                fontSize=12,
                spaceBefore=8,
                spaceAfter=2,
                textColor=colors.HexColor('#1F2937'),
                fontName='Helvetica-Bold'
            )
            
            company_style = ParagraphStyle(
                'CompanyStyle',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=2,
                textColor=colors.HexColor(self.uae_colors['primary']),
                fontName='Helvetica'
            )
            
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                textColor=colors.HexColor('#6B7280'),
                fontName='Helvetica'
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                textColor=colors.HexColor('#374151'),
                fontName='Helvetica',
                leading=14
            )
            
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=3,
                textColor=colors.HexColor('#374151'),
                fontName='Helvetica',
                leftIndent=15,
                leading=13
            )
            
            # Add content
            data = cv_data['data']
            personal_info = get_field(data, 'personal_info', 'personalInfo', default={})
            
            # Header with name
            full_name = get_field(personal_info, 'full_name', 'fullName', 'firstName')
            if not full_name:
                # Try combining first and last name
                first = get_field(personal_info, 'firstName', 'first_name', default='')
                last = get_field(personal_info, 'lastName', 'last_name', default='')
                full_name = f"{first} {last}".strip()
            
            if full_name:
                story.append(Paragraph(full_name, name_style))
            
            # Contact information
            contact_parts = []
            email = get_field(personal_info, 'email')
            phone = get_field(personal_info, 'phone')
            location = get_field(personal_info, 'location')
            if not location:
                city = get_field(personal_info, 'city')
                emirate = get_field(personal_info, 'emirate')
                country = get_field(personal_info, 'country', default='United Arab Emirates')
                if city or emirate:
                    location = ', '.join(filter(None, [city, emirate, country]))
                elif country:
                    location = country
            
            if email:
                contact_parts.append(f"✉ {email}")
            if phone:
                contact_parts.append(f"☎ {phone}")
            if location:
                contact_parts.append(f"📍 {location}")
            
            if contact_parts:
                story.append(Paragraph(" | ".join(contact_parts), contact_style))
            
            # Professional Summary
            summary = get_field(data, 'professional_summary', 'professionalSummary')
            if summary:
                story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
                story.append(Paragraph(summary, body_style))
            
            # Core Competencies / Skills (show before experience like in preview)
            skills = get_field(data, 'skills', default=[])
            if skills:
                story.append(Paragraph("CORE COMPETENCIES", heading_style))
                skills_by_category = {}
                
                for skill in skills:
                    if isinstance(skill, str):
                        category = 'Skills'
                        skill_name = skill
                    elif isinstance(skill, dict):
                        category = get_field(skill, 'category', default='Skills')
                        skill_name = get_field(skill, 'name', default=str(skill))
                    else:
                        category = 'Skills'
                        skill_name = str(skill)
                    
                    if category not in skills_by_category:
                        skills_by_category[category] = []
                    skills_by_category[category].append(skill_name)
                
                for category, skill_list in skills_by_category.items():
                    skills_text = f"<b>{category}:</b> {', '.join(skill_list)}"
                    story.append(Paragraph(skills_text, body_style))
                
                story.append(Spacer(1, 10))
            
            # Professional Experience
            experience = get_field(data, 'experience', default=[])
            if experience:
                story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
                
                for exp in experience:
                    if not isinstance(exp, dict):
                        continue
                    
                    # Job title - try multiple field names
                    job_title = get_field(exp, 'jobTitle', 'job_title', 'position', 'title', default='')
                    company = get_field(exp, 'company', 'employer_admin', 'organization', default='')
                    
                    # Dates
                    start_date = get_field(exp, 'startDate', 'start_date', default='')
                    end_date = get_field(exp, 'endDate', 'end_date', default='')
                    is_current = get_field(exp, 'isCurrentJob', 'isCurrentlyWorking', 'is_current', 'isCurrent', default=False)
                    
                    if is_current or (isinstance(end_date, str) and end_date.lower() in ['present', 'current', '']):
                        end_date = 'Present'
                    
                    location = get_field(exp, 'location', default='')
                    
                    # Build the experience entry
                    if job_title:
                        story.append(Paragraph(job_title.upper(), job_title_style))
                    
                    if company:
                        story.append(Paragraph(company, company_style))
                    
                    date_parts = []
                    if start_date and end_date:
                        date_parts.append(f"{start_date} - {end_date}")
                    elif start_date:
                        date_parts.append(f"{start_date} - Present")
                    if location:
                        date_parts.append(location)
                    
                    if date_parts:
                        story.append(Paragraph(" | ".join(date_parts), date_style))
                    
                    # Description - handle both string and list formats
                    description = get_field(exp, 'description', 'responsibilities', default='')
                    if description:
                        if isinstance(description, list):
                            for desc in description:
                                story.append(Paragraph(f"• {desc}", bullet_style))
                        elif isinstance(description, str) and description.strip():
                            # Split by newlines or periods for bullet points
                            story.append(Paragraph(f"• {description}", bullet_style))
                    
                    # Achievements
                    achievements = get_field(exp, 'achievements', 'accomplishments', default=[])
                    if achievements and isinstance(achievements, list):
                        for achievement in achievements:
                            story.append(Paragraph(f"• {achievement}", bullet_style))
                    
                    story.append(Spacer(1, 8))
            
            # Education
            education = get_field(data, 'education', default=[])
            if education:
                story.append(Paragraph("EDUCATION", heading_style))
                
                for edu in education:
                    if not isinstance(edu, dict):
                        continue
                    
                    degree = get_field(edu, 'degree', 'qualification', 'title', default='')
                    institution = get_field(edu, 'institution', 'school', 'university', default='')
                    field = get_field(edu, 'fieldOfStudy', 'field_of_study', 'major', default='')
                    location = get_field(edu, 'location', default='')
                    start_date = get_field(edu, 'startDate', 'start_date', default='')
                    end_date = get_field(edu, 'endDate', 'end_date', 'graduationYear', 'graduation_year', default='')
                    gpa = get_field(edu, 'gpa', 'grade', default='')
                    
                    # Build education entry
                    if degree:
                        story.append(Paragraph(f"<b>{degree}</b>", body_style))
                    
                    inst_parts = [institution]
                    if location:
                        inst_parts.append(location)
                    if end_date:
                        inst_parts.append(str(end_date))
                    
                    if institution:
                        story.append(Paragraph(" • ".join(filter(None, inst_parts)), body_style))
                    
                    if field:
                        story.append(Paragraph(f"Field of Study: {field}", body_style))
                    
                    if gpa:
                        story.append(Paragraph(f"GPA: {gpa}", body_style))
                    
                    story.append(Spacer(1, 6))
            
            # Languages
            languages = get_field(data, 'languages', default=[])
            if languages:
                story.append(Paragraph("LANGUAGES", heading_style))
                for lang in languages:
                    if isinstance(lang, dict):
                        lang_name = get_field(lang, 'name', 'language', default='')
                        proficiency = get_field(lang, 'proficiency', 'level', default='')
                        if lang_name:
                            lang_text = f"{lang_name}"
                            if proficiency:
                                lang_text += f" - {proficiency}"
                            story.append(Paragraph(lang_text, body_style))
                    elif isinstance(lang, str):
                        story.append(Paragraph(lang, body_style))
                story.append(Spacer(1, 10))
            
            # Certifications
            certifications = get_field(data, 'certifications', 'certificates', default=[])
            if certifications:
                story.append(Paragraph("CERTIFICATIONS", heading_style))
                for cert in certifications:
                    if isinstance(cert, dict):
                        cert_name = get_field(cert, 'name', 'title', default='')
                        issuer = get_field(cert, 'issuer', 'organization', 'issuingOrganization', default='')
                        date = get_field(cert, 'date', 'issueDate', 'issue_date', default='')
                        if cert_name:
                            cert_text = f"<b>{cert_name}</b>"
                            if issuer:
                                cert_text += f" - {issuer}"
                            if date:
                                cert_text += f" ({date})"
                            story.append(Paragraph(cert_text, body_style))
                    elif isinstance(cert, str):
                        story.append(Paragraph(cert, body_style))
                story.append(Spacer(1, 10))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Exported CV {cv_id} to PDF: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error exporting PDF: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
    
    def _export_simple_pdf(self, cv_data: Dict[str, Any], cv_id: str, timestamp: str) -> str:
        """Fallback simple PDF export when ReportLab is not available"""
        # Create a simple text file as fallback
        filename = f"cv_{cv_id}_{timestamp}.txt"
        file_path = self.exports_dir / filename
        
        data = cv_data['data']
        personal_info = get_field(data, 'personal_info', 'personalInfo', default={})
        
        lines = []
        
        # Name
        full_name = get_field(personal_info, 'full_name', 'fullName', default='')
        if full_name:
            lines.append(full_name.upper())
            lines.append("=" * len(full_name))
        
        # Contact
        email = get_field(personal_info, 'email')
        phone = get_field(personal_info, 'phone')
        if email:
            lines.append(f"Email: {email}")
        if phone:
            lines.append(f"Phone: {phone}")
        
        lines.append("")
        
        # Summary
        summary = get_field(data, 'professional_summary', 'professionalSummary')
        if summary:
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 20)
            lines.append(summary)
            lines.append("")
        
        # Experience
        experience = get_field(data, 'experience', default=[])
        if experience:
            lines.append("WORK EXPERIENCE")
            lines.append("-" * 20)
            for exp in experience:
                if isinstance(exp, dict):
                    job_title = get_field(exp, 'jobTitle', 'job_title', 'position', default='')
                    company = get_field(exp, 'company', default='')
                    lines.append(f"{job_title} at {company}")
            lines.append("")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(lines))
        
        return str(file_path)
    
    def _export_docx(self, cv_data: Dict[str, Any], cv_id: str, timestamp: str) -> str:
        """Export CV as DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export")
        
        try:
            filename = f"cv_{cv_id}_{timestamp}.docx"
            file_path = self.exports_dir / filename
            
            doc = Document()
            data = cv_data['data']
            personal_info = get_field(data, 'personal_info', 'personalInfo', default={})
            
            # Name
            full_name = get_field(personal_info, 'full_name', 'fullName', default='')
            if not full_name:
                first = get_field(personal_info, 'firstName', 'first_name', default='')
                last = get_field(personal_info, 'lastName', 'last_name', default='')
                full_name = f"{first} {last}".strip()
            
            if full_name:
                heading = doc.add_heading(full_name, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact info
            contact_parts = []
            email = get_field(personal_info, 'email')
            phone = get_field(personal_info, 'phone')
            location = get_field(personal_info, 'location')
            
            if email:
                contact_parts.append(email)
            if phone:
                contact_parts.append(phone)
            if location:
                contact_parts.append(location)
            
            if contact_parts:
                contact_para = doc.add_paragraph(" | ".join(contact_parts))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Professional Summary
            summary = get_field(data, 'professional_summary', 'professionalSummary')
            if summary:
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(summary)
            
            # Experience
            experience = get_field(data, 'experience', default=[])
            if experience:
                doc.add_heading('Professional Experience', level=1)
                for exp in experience:
                    if not isinstance(exp, dict):
                        continue
                    
                    job_title = get_field(exp, 'jobTitle', 'job_title', 'position', default='')
                    company = get_field(exp, 'company', default='')
                    start_date = get_field(exp, 'startDate', 'start_date', default='')
                    end_date = get_field(exp, 'endDate', 'end_date', default='')
                    is_current = get_field(exp, 'isCurrentJob', 'isCurrentlyWorking', 'is_current', default=False)
                    
                    if is_current:
                        end_date = 'Present'
                    
                    # Job title
                    if job_title:
                        p = doc.add_paragraph()
                        p.add_run(job_title).bold = True
                    
                    # Company and dates
                    if company:
                        date_str = f"{start_date} - {end_date}" if start_date else ""
                        doc.add_paragraph(f"{company} | {date_str}")
                    
                    # Description
                    description = get_field(exp, 'description', default='')
                    if description:
                        if isinstance(description, list):
                            for desc in description:
                                doc.add_paragraph(f"• {desc}")
                        else:
                            doc.add_paragraph(f"• {description}")
            
            # Education
            education = get_field(data, 'education', default=[])
            if education:
                doc.add_heading('Education', level=1)
                for edu in education:
                    if not isinstance(edu, dict):
                        continue
                    
                    degree = get_field(edu, 'degree', default='')
                    institution = get_field(edu, 'institution', default='')
                    
                    if degree:
                        p = doc.add_paragraph()
                        p.add_run(degree).bold = True
                    
                    if institution:
                        doc.add_paragraph(institution)
            
            # Skills
            skills = get_field(data, 'skills', default=[])
            if skills:
                doc.add_heading('Skills', level=1)
                skills_by_category = {}
                
                for skill in skills:
                    if isinstance(skill, str):
                        category = 'Skills'
                        skill_name = skill
                    elif isinstance(skill, dict):
                        category = get_field(skill, 'category', default='Skills')
                        skill_name = get_field(skill, 'name', default=str(skill))
                    else:
                        category = 'Skills'
                        skill_name = str(skill)
                    
                    if category not in skills_by_category:
                        skills_by_category[category] = []
                    skills_by_category[category].append(skill_name)
                
                for category, skill_list in skills_by_category.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{category}: ").bold = True
                    p.add_run(", ".join(skill_list))
            
            doc.save(str(file_path))
            
            logger.info(f"Exported CV {cv_id} to DOCX: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"Error exporting DOCX: {str(e)}")
            raise
