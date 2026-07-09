#!/usr/bin/env python3
"""
CV Parser with Qwen / DashScope Integration
Emirati Journey Platform — Qwen Migration
"""

import os
import json
import logging
import uuid
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
import tempfile
import mimetypes
from pathlib import Path

from werkzeug.datastructures import FileStorage

# Qwen / DashScope client (replaces google.generativeai)
try:
    from backend.services.qwen_client import chat_completion, QwenParsingError, QwenClientError
    from backend.config.qwen_config import DASHSCOPE_API_KEY, MAX_INPUT_CHARS
    _qwen_available = bool(DASHSCOPE_API_KEY)
except ImportError:
    _qwen_available = False

# PDF / DOCX extraction
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import docx
    from docx import Document
except ImportError:
    docx = None
    Document = None

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class CVParser:
    """CV Parser with Qwen / DashScope integration for comprehensive CV analysis"""
    
    def __init__(self):
        """Initialize CV Parser with Qwen configuration"""
        self.ai_available = _qwen_available
        
        if not self.ai_available:
            logger.warning("⚠️ DASHSCOPE_API_KEY not set. AI features will be disabled.")
        else:
            logger.info("✅ CV Parser initialized with Qwen / DashScope (qwen-turbo)")

        
        # Supported file types
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt'
        }
        
        # Create temp directory for file processing
        self.temp_dir = Path(tempfile.gettempdir()) / 'cv_parser'
        self.temp_dir.mkdir(exist_ok=True)

    def parse_cv(self, file_path: str, user_id: str = None) -> Dict[str, Any]:
        """
        Parse CV from a file path (Wrapper for compatibility).
        This is the method called by enhanced_cv_routes.py
        """
        try:
            if not os.path.exists(file_path):
                 return {'success': False, 'message': f'File not found: {file_path}'}
            
            # Detect mime type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                 # Fallback based on extension
                 ext = os.path.splitext(file_path)[1].lower()
                 if ext == '.pdf': mime_type = 'application/pdf'
                 elif ext == '.docx': mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                 elif ext == '.txt': mime_type = 'text/plain'
            
            # Extract text
            text_content = ""
            if mime_type == 'application/pdf':
                with open(file_path, 'rb') as f:
                     text_content = self._extract_from_pdf_stream(f)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                 # docx library handles paths directly usually, but let's stick to stream for consistency if we want
                 # But _extract_from_docx uses Document(path) in my previous read?
                 # Let's write a robust extractor for path
                 text_content = self._extract_text_from_path(file_path, mime_type)
            elif mime_type == 'text/plain':
                 with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                     text_content = f.read()
            
            if not text_content:
                 return {'success': False, 'message': 'Could not extract text from file'}

            return self.parse_cv_text(text_content, user_id, os.path.basename(file_path))

        except Exception as e:
            logger.error(f"Error in parse_cv: {e}")
            return {'success': False, 'message': str(e)}

    
    def parse_cv_file(self, file: FileStorage, user_id: str = None) -> Dict[str, Any]:
        """Parse CV from uploaded FileStorage object"""
        try:
            # Validate file
            if not file or not file.filename:
                return {
                    'success': False,
                    'message': 'No file provided'
                }
            
            # Check file type
            mime_type = file.content_type or mimetypes.guess_type(file.filename)[0]
            # Simple check
            if mime_type not in self.supported_types and not any(file.filename.lower().endswith(ext) for ext in ['.pdf', '.docx', '.doc', '.txt']):
                return {
                    'success': False,
                    'message': f'Unsupported file type: {mime_type}. Supported: PDF, DOCX, DOC, TXT'
                }
            
            # Extract text from file
            text_content = self._extract_text_from_file_storage(file, mime_type)
            if not text_content:
                return {
                    'success': False,
                    'message': 'Could not extract text from file'
                }
            
            # Parse the extracted text
            return self.parse_cv_text(text_content, user_id, file.filename)
            
        except Exception as e:
            logger.error(f"Error parsing CV file: {str(e)}")
            return {
                'success': False,
                'message': f'File parsing failed: {str(e)}'
            }
    
    def parse_cv_text(self, text: str, user_id: str = None, filename: str = None) -> Dict[str, Any]:
        """Parse CV from text content using Qwen via DashScope"""
        try:
            cleaned_text = text.strip()
            if not cleaned_text or len(cleaned_text) < 50:
                return {
                    'success': False,
                    'message': 'Text content too short or empty'
                }

            # If no API key, use fallback
            if not self.ai_available:
                 logger.info("Using fallback parser (No AI)")
                 return self._fallback_parse_text(cleaned_text, user_id, filename)
            
            # Generate CV ID
            cv_id = str(uuid.uuid4())
            
            # Build chat messages for Qwen
            import time
            start_time = time.time()

            messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an expert AI Resume Parser built for the UAE job market "
                        "and Emirati workforce. You MUST return ONLY raw, valid JSON. "
                        "No markdown, no code fences, no explanatory text. "
                        "Field keys MUST be in English. Values can be in Arabic or English "
                        "per the source CV."
                    ),
                },
                {
                    "role": "user",
                    "content": self._create_parsing_prompt(cleaned_text),
                },
            ]

            try:
                parsed_data = chat_completion(
                    task_type="parse",
                    messages=messages,
                    response_format={"type": "json_object"},
                )
            except (QwenParsingError, QwenClientError) as e:
                logger.error(f"❌ Qwen CV parsing failed: {e}")
                return self._fallback_parse_text(cleaned_text, user_id, filename)

            processing_time = round(time.time() - start_time, 3)

            if not parsed_data:
                logger.warning("Empty response from Qwen, using fallback")
                return self._fallback_parse_text(cleaned_text, user_id, filename)
            
            # Enhance with UAE-specific analysis
            enhanced_data = self._enhance_with_uae_analysis(parsed_data, cleaned_text)
            
            # Calculate scores and insights
            analysis_results = self._calculate_cv_scores(enhanced_data, cleaned_text)
            
            # Prepare final result
            result = {
                'success': True,
                'cv_id': cv_id,
                'data': {
                    'personal_info': enhanced_data.get('personal_info', {}),
                    'professional_summary': enhanced_data.get('professional_summary', ''),
                    'experience': enhanced_data.get('experience', []),
                    'education': enhanced_data.get('education', []),
                    'skills': enhanced_data.get('skills', []),
                    'languages': enhanced_data.get('languages', []),
                    'certifications': enhanced_data.get('certifications', []),
                    'achievements': enhanced_data.get('achievements', []),
                    'projects': enhanced_data.get('projects', []),
                    'references': enhanced_data.get('references', [])
                },
                'analysis': analysis_results,
                'metadata': {
                    'cv_id': cv_id,
                    'user_id': user_id,
                    'filename': filename,
                    'parsed_at': datetime.utcnow().isoformat(),
                    'parser_version': '3.0-qwen',
                    'ai_model': 'qwen-turbo',
                    'text_length': len(cleaned_text),
                    'processing_time': processing_time
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing CV text with Qwen: {str(e)}")
            return self._fallback_parse_text(text, user_id, filename)

    def _fallback_parse_text(self, text: str, user_id: str = None, filename: str = None) -> Dict[str, Any]:
        """Fallback parsing when AI is unavailable"""
        cv_id = str(uuid.uuid4())
        import re
        
        # Simple RegExp Extraction
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
        email = email_match.group(0) if email_match else ""
        
        phone_match = re.search(r'(\+971|05)\d{8,9}', text)
        phone = phone_match.group(0) if phone_match else ""
        
        lines = [l.strip() for l in text.split('\n') if l.strip()]  # [FIX] Use actual newline, not literal backslash-n
        name = lines[0][:100] if lines else "Candidate"  # [FIX] Truncate name to prevent DB errors
        
        name_parts = name.split()
        first_name = name_parts[0] if name_parts else ""
        last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

        data = {
            'personal_info': {
                'full_name': name,
                'first_name': first_name,
                'last_name': last_name,
                'email': email,
                'phone': phone,
                'nationality': 'UAE' 
            },
            # [FIX] Truncate professional_summary to prevent DB errors (keep first 200 chars max)
            'professional_summary': text[:200] if len(text) > 200 else text,
            # [FIX] Extract basic experience, education, skills from text patterns
            'experience': self._extract_basic_experience(text),
            'education': self._extract_basic_education(text),
            'skills': self._extract_basic_skills(text),
            'languages': [],
            'certifications': [],
        }
        
        analysis_results = {
            'scores': {'overall': 40, 'completeness': 40, 'uae_relevance': 50}, 
            'insights': {'message': 'Parsed using basic fallback parser (AI unavailable)'},
            'recommendations': ['Please review and update your details manually.']
        }

        result = {
            'success': True,
            'cv_id': cv_id,
            'data': data,
            'analysis': analysis_results,
            'metadata': {
                'cv_id': cv_id,
                'user_id': user_id,
                'filename': filename,
                'parsed_at': datetime.utcnow().isoformat(),
                'parser_version': '2.0-fallback',
                'ai_model': 'none',
                'text_length': len(text)
            }
        }
        return result
    
    def _extract_basic_experience(self, text: str) -> list:
        """Extract basic work experience using pattern matching"""
        experience = []
        lines = text.split('\n')
        
        # Look for common job title keywords
        job_keywords = ['engineer', 'developer', 'manager', 'analyst', 'consultant', 
                       'director', 'coordinator', 'specialist', 'assistant', 'officer',
                       'lead', 'senior', 'junior', 'intern', 'executive']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            # Check if line contains a job title
            if any(keyword in line_lower for keyword in job_keywords) and len(line.strip()) > 5:
                # Try to extract company from next few lines
                company = "Company"
                for j in range(i+1, min(i+3, len(lines))):
                    if lines[j].strip() and len(lines[j].strip()) < 50:
                        company = lines[j].strip()[:100]
                        break
                
                experience.append({
                    'company': company,
                    'position': line.strip()[:100],
                    'start_date': None,
                    'end_date': None,
                    'is_current': False,
                    'description': ''
                })
                
                # Limit to 3 entries to avoid spam
                if len(experience) >= 3:
                    break
        
        return experience if experience else []
    
    def _extract_basic_education(self, text: str) -> list:
        """Extract basic education using pattern matching"""
        education = []
        lines = text.split('\n')
        
        # Look for degree keywords
        degree_keywords = ['bachelor', 'master', 'phd', 'diploma', 'degree', 
                          'university', 'college', 'institute', 'school',
                          'b.sc', 'm.sc', 'bsc', 'msc', 'ba', 'ma', 'mba']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in degree_keywords) and len(line.strip()) > 5:
                # Extract institution and degree
                degree = line.strip()[:100]
                institution = "Educational Institution"
                
                # Try to find institution in nearby lines
                for j in range(max(0, i-2), min(i+3, len(lines))):
                    if 'university' in lines[j].lower() or 'college' in lines[j].lower():
                        institution = lines[j].strip()[:100]
                        break
                
                education.append({
                    'institution': institution,
                    'degree': degree,
                    'field_of_study': '',
                    'start_date': None,
                    'end_date': None,
                    'grade': ''
                })
                
                # Limit to 2 entries
                if len(education) >= 2:
                    break
        
        return education if education else []
    
    def _extract_basic_skills(self, text: str) -> list:
        """Extract basic skills using common skill keywords"""
        skills = []
        text_lower = text.lower()
        
        # Common technical skills
        tech_skills = ['python', 'java', 'javascript', 'react', 'node.js', 'sql', 
                      'c++', 'c#', 'php', 'ruby', 'go', 'swift', 'kotlin',
                      'html', 'css', 'typescript', 'angular', 'vue',
                      'aws', 'azure', 'docker', 'kubernetes', 'git']
        
        # Common soft skills
        soft_skills = ['leadership', 'communication', 'teamwork', 'problem solving',
                      'analytical', 'creative', 'organized', 'detail-oriented']
        
        # Extract technical skills
        for skill in tech_skills:
            if skill in text_lower:
                skills.append({
                    'name': skill.title(),
                    'level': 'Intermediate',
                    'category': 'Technical'
                })
        
        # Extract soft skills  
        for skill in soft_skills:
            if skill in text_lower:
                skills.append({
                    'name': skill.title(),
                    'level': 'Intermediate',
                    'category': 'Soft Skill'
                })
        
        # Limit to 10 skills
        return skills[:10] if skills else [{'name': 'General', 'level': 'Basic', 'category': 'General'}]
    
    
    def _extract_text_from_file_storage(self, file: FileStorage, mime_type: str) -> str:
        """Extract text content from uploaded file"""
        try:
            # We must map mime types to our logic
            if 'pdf' in mime_type:
                return self._extract_from_pdf_stream(file)
            elif 'word' in mime_type or 'office' in mime_type:
                 # Need to save to temp for docx usually
                suffix = '.docx' if 'openxml' in mime_type else '.doc'
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    file.save(tmp.name)
                    tmp_path = tmp.name
                
                text = self._extract_text_from_path(tmp_path, mime_type)
                
                try:
                    os.unlink(tmp_path)
                except: pass
                return text

            elif 'text' in mime_type:
                return file.read().decode('utf-8', errors='ignore')
            else:
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}")
            return ""

    def _extract_from_pdf_stream(self, file_stream) -> str:
        try:
            if pdfplumber:
                # pdfplumber needs a file path or file-like object
                file_stream.seek(0)
                with pdfplumber.open(file_stream) as pdf:
                    pages_text = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            pages_text.append(text)
                        # Also extract tables for comprehensive data capture
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = ' | '.join([cell or '' for cell in row])
                                    if row_text.strip():
                                        pages_text.append(row_text)
                    return '\n'.join(pages_text).strip()
            else:
                logger.warning("pdfplumber not available, returning empty text")
                return ""
        except Exception as e:
            logger.error(f"PDF extract error: {e}")
            return ""

    def _extract_text_from_path(self, path: str, mime_type: str) -> str:
        try:
            if 'pdf' in mime_type or path.lower().endswith('.pdf'):
                with open(path, 'rb') as f:
                    return self._extract_from_pdf_stream(f)
            elif ('word' in mime_type or 'office' in mime_type or path.endswith('.doc') or path.endswith('.docx')):
                 if docx:
                     doc = Document(path)
                     text = "\n".join([p.text for p in doc.paragraphs])
                     for table in doc.tables:
                         for row in table.rows:
                             text += " ".join([c.text for c in row.cells]) + "\n"
                     return text
                 else:
                     return ""
            else:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Path extract error: {e}")
            return ""

    def _create_parsing_prompt(self, text: str) -> str:
        return f"""You are an expert AI Resume Parser built for the UAE job market and Emirati workforce.
Extract structured JSON from this CV/resume. Follow every rule precisely.

═══════════════════════════════════════════
LANGUAGE & ENCODING
═══════════════════════════════════════════
- The CV may be in English, Arabic, or bilingual. Extract ALL content regardless of language.
- For Arabic names, transliterate to Latin script in `full_name` AND keep the original in `full_name_ar`.
- If the CV is entirely in Arabic, still return field keys in English with Arabic values.

═══════════════════════════════════════════
EXTRACTION RULES
═══════════════════════════════════════════
1. **Name**: Almost always the very first line. Extract `full_name` strictly. Split into `first_name` and `last_name`.
2. **Dates**: MUST be "YYYY-MM-DD" or null.
   - Currently working → `end_date`: null, `is_current`: true. NEVER use "Present" or "Current".
   - Day unknown → use "01". Month unknown → use "01". (e.g., "2022" → "2022-01-01")
   - Durations like "5 years" → do NOT guess dates. Leave as null and note in description.
3. **Summary**: Maximum 2 sentences, 30 words. Be extremely concise.
4. **Skills**: Extract ALL technical and soft skills. Categorize each as "Technical", "Soft", "Language", or "Domain".
   - Include proficiency level: "Beginner", "Intermediate", "Advanced", or "Expert".
5. **Experience**: Extract company, position, dates, location, and a description of responsibilities.
   - Also extract key `achievements` as a separate array of strings per role.
6. **Education**: Extract institution, degree, field_of_study, dates, and GPA if mentioned.
7. **Phone**: UAE numbers use formats like +971-50-XXX-XXXX or 05XXXXXXXX. Normalize to include country code.
8. **Projects**: Extract any personal, academic, or professional projects mentioned.
9. **Certifications**: Include issuer and date obtained if available.

═══════════════════════════════════════════
UAE NATIONAL QUALIFICATIONS FRAMEWORK (NQF)
═══════════════════════════════════════════
Map each education entry to the UAE NQF level. Use this mapping:
- Level 1: Certificate — basic vocational
- Level 2: Certificate — skilled vocational
- Level 3: Certificate — advanced vocational / High School Diploma
- Level 4: Certificate — post-secondary vocational
- Level 5: Diploma / Associate Degree
- Level 6: Advanced Diploma / Higher Diploma
- Level 7: Bachelor's Degree
- Level 8: Postgraduate Diploma / Bachelor's Honours
- Level 9: Master's Degree / MBA / EMBA
- Level 10: Doctoral Degree (PhD / DBA)

Set `nqf_level` as an integer (1-10) on each education entry. If unclear, estimate from the degree name.
Also set `highest_nqf_level` at the top level as the maximum across all education entries.

═══════════════════════════════════════════
CV TEXT
═══════════════════════════════════════════
{text[:20000]}

═══════════════════════════════════════════
OUTPUT (return ONLY this JSON, no markdown, no explanation)
═══════════════════════════════════════════
{{
    "personal_info": {{
        "full_name": "Abdulaziz Al Falahi",
        "full_name_ar": "عبدالعزيز الفلاحي",
        "first_name": "Abdulaziz",
        "last_name": "Al Falahi",
        "email": "user@example.com",
        "phone": "+971501234567",
        "nationality": "UAE National",
        "location": "Dubai, UAE",
        "address": "Dubai, UAE",
        "linkedin": ""
    }},
    "professional_summary": "Concise 2-sentence summary here.",
    "highest_nqf_level": 7,
    "total_experience_years": 5,
    "skills": [
        {{ "name": "Python", "level": "Advanced", "category": "Technical" }},
        {{ "name": "Leadership", "level": "Expert", "category": "Soft" }},
        {{ "name": "Arabic", "level": "Native", "category": "Language" }}
    ],
    "experience": [
        {{
            "company": "Emirates NBD",
            "position": "Senior Developer",
            "start_date": "2020-01-01",
            "end_date": null,
            "is_current": true,
            "location": "Dubai, UAE",
            "description": "Led development of digital banking platform.",
            "achievements": ["Increased platform uptime to 99.9%", "Managed team of 8 engineers"]
        }}
    ],
    "education": [
        {{
            "institution": "Khalifa University",
            "degree": "Bachelor of Science",
            "field_of_study": "Computer Science",
            "start_date": "2015-09-01",
            "end_date": "2019-06-01",
            "gpa": "3.8",
            "nqf_level": 7
        }}
    ],
    "certifications": [
        {{ "name": "AWS Solutions Architect", "issuer": "Amazon", "date": "2023-03-01" }}
    ],
    "languages": [
        {{ "language": "Arabic", "proficiency": "Native" }},
        {{ "language": "English", "proficiency": "Fluent" }}
    ],
    "projects": [
        {{ "name": "Smart City Dashboard", "description": "Real-time analytics platform", "technologies": ["React", "Python", "AWS"] }}
    ],
    "volunteer_work": []
}}"""

    # NOTE: _parse_gemini_response removed — JSON extraction is handled by
    # qwen_client._extract_json() automatically within chat_completion().

    def _enhance_with_uae_analysis(self, data: Dict, text: str) -> Dict:
        # Placeholder for UAE Specific Logic
        if 'uae_analysis' not in data:
            data['uae_analysis'] = {
                'is_uae_national': 'uae' in text.lower() or 'emirati' in text.lower(),
                'uae_experience_years': 0,
                'has_arabic_language': False
            }
        return data

    def _calculate_cv_scores(self, data: Dict, text: str) -> Dict:
        """Calculate real CV quality scores using weighted criteria"""
        scores = {}
        recommendations = []
        
        # --- Completeness Score (40% weight) ---
        completeness_fields = {
            'name': bool(data.get('personal_info', {}).get('full_name')),
            'email': bool(data.get('personal_info', {}).get('email')),
            'phone': bool(data.get('personal_info', {}).get('phone')),
            'summary': bool(data.get('professional_summary')),
            'experience': len(data.get('experience', [])) > 0,
            'education': len(data.get('education', [])) > 0,
            'skills': len(data.get('skills', [])) > 0,
            'certifications': len(data.get('certifications', [])) > 0,
        }
        filled = sum(1 for v in completeness_fields.values() if v)
        completeness = int((filled / len(completeness_fields)) * 100)
        scores['completeness'] = completeness
        
        missing = [k for k, v in completeness_fields.items() if not v]
        if missing:
            recommendations.append(f"Add missing sections: {', '.join(missing)}")
        
        # --- Detail Depth Score (30% weight) ---
        depth_points = 0
        max_depth = 100
        
        # Experience detail
        experiences = data.get('experience', [])
        if len(experiences) >= 3:
            depth_points += 25
        elif len(experiences) >= 1:
            depth_points += 15
        else:
            recommendations.append('Add work experience entries with descriptions')
        
        # Check experience descriptions
        has_descriptions = sum(1 for e in experiences if len(e.get('description', '')) > 20)
        if has_descriptions >= 2:
            depth_points += 20
        elif has_descriptions >= 1:
            depth_points += 10
        else:
            recommendations.append('Add detailed descriptions to your work experience')
        
        # Skills count
        skills_count = len(data.get('skills', []))
        if skills_count >= 8:
            depth_points += 25
        elif skills_count >= 4:
            depth_points += 15
        elif skills_count >= 1:
            depth_points += 8
        else:
            recommendations.append('List your technical and soft skills')
        
        # Education detail
        education = data.get('education', [])
        if len(education) >= 1:
            depth_points += 15
            if any(e.get('field_of_study') for e in education):
                depth_points += 10
        
        # Summary length
        summary = data.get('professional_summary', '')
        if len(summary) > 100:
            depth_points += 5
        elif len(summary) > 30:
            depth_points += 3
        
        depth_score = min(int(depth_points), max_depth)
        scores['detail_depth'] = depth_score
        
        # --- UAE Relevance Score (30% weight) ---
        uae_points = 0
        text_lower = text.lower()
        
        # Nationality check
        uae_keywords = ['uae', 'emirati', 'united arab emirates', 'dubai', 'abu dhabi', 
                       'sharjah', 'ajman', 'fujairah', 'ras al khaimah', 'umm al quwain']
        if any(kw in text_lower for kw in uae_keywords):
            uae_points += 30
        
        # Arabic language
        if 'arabic' in text_lower or 'عربي' in text:
            uae_points += 20
        else:
            recommendations.append('Mention Arabic language proficiency if applicable')
        
        # UAE companies/institutions
        uae_entities = ['adnoc', 'etisalat', 'emaar', 'dubai government', 'mubadala',
                       'masdar', 'emirates', 'du telecom', 'dewa', 'rta',
                       'american university of sharjah', 'khalifa university',
                       'zayed university', 'university of sharjah']
        entity_matches = sum(1 for e in uae_entities if e in text_lower)
        uae_points += min(entity_matches * 10, 30)
        
        # Emiratization keywords
        emiratization_kw = ['emiratization', 'nafis', 'national service', 'tawteen']
        if any(kw in text_lower for kw in emiratization_kw):
            uae_points += 20
        
        uae_relevance = min(uae_points, 100)
        scores['uae_relevance'] = uae_relevance
        
        # --- Overall Score ---
        overall = int(completeness * 0.4 + depth_score * 0.3 + uae_relevance * 0.3)
        scores['overall'] = min(overall, 100)
        
        # Insights
        insights = {
            'strengths': [],
            'improvement_areas': []
        }
        if completeness >= 80:
            insights['strengths'].append('Comprehensive CV with all key sections')
        if depth_score >= 70:
            insights['strengths'].append('Detailed experience and skills descriptions')
        if uae_relevance >= 60:
            insights['strengths'].append('Strong UAE market relevance')
        if completeness < 60:
            insights['improvement_areas'].append('CV is missing several important sections')
        if depth_score < 50:
            insights['improvement_areas'].append('Add more detail to experience and skills')
        if uae_relevance < 30:
            insights['improvement_areas'].append('Consider adding UAE-specific context')
        
        return {'scores': scores, 'insights': insights, 'recommendations': recommendations}

# Instantiate the parser for import
cv_parser = CVParser()
