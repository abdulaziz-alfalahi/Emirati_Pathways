"""A table in a CV must reach the model with its columns intact.

WHY THIS FILE EXISTS

`cv_parser._extract_text_from_path` joined table cells with a plain space, so a
document's tables arrived as a single run of words:

    Acme Corp Senior Analyst 2019 2023

From that, nothing can tell the employer from the job title from the dates —
and employment history and education are exactly the things CVs put in tables.
The PDF branch a few lines above, and services/pdf_extractor.py, both already
used ' | '. This one path had simply not followed the convention.

The same code also appended the tables to a string with `+=` and no separator,
so the last paragraph of the document ran into the first table row
("…الإماراتية، دبيHost Certificates sent"), fusing an unrelated sentence to a
table header.

Found on 2026-08-24 while comparing our extraction against markitdown. We chose
not to adopt markitdown — measured, it was indistinguishable on PDFs and
sometimes worse on HTML — but its DOCX table handling exposed this.
"""
import os
import sys

BACKEND = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BACKEND)

import pytest  # noqa: E402

docx = pytest.importorskip('docx', reason='python-docx not installed')
from docx import Document  # noqa: E402

DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def _doc_with_table(path):
    d = Document()
    d.add_paragraph('Curriculum Vitae')
    d.add_paragraph('Experience follows.')
    t = d.add_table(rows=3, cols=3)
    rows = [('Employer', 'Role', 'Dates'),
            ('Acme Corp', 'Senior Analyst', '2019-2023'),
            ('Beta LLC', 'Analyst', '2016-2019')]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    d.save(path)
    return path


def _extract(path):
    from cv_parser import CVParser
    return CVParser()._extract_text_from_path(path, DOCX_MIME)


def test_table_columns_survive_extraction(tmp_path):
    text = _extract(_doc_with_table(str(tmp_path / 'cv.docx')))
    assert 'Acme Corp | Senior Analyst | 2019-2023' in text, (
        'table cells are not separated, so employer, role and dates arrive '
        f'fused into one run. Got:\n{text}'
    )


def test_each_table_row_is_its_own_line(tmp_path):
    """Two jobs on one line would read as one job with four dates."""
    text = _extract(_doc_with_table(str(tmp_path / 'cv.docx')))
    lines = [l for l in text.splitlines() if ' | ' in l]
    assert len(lines) == 3, f'expected 3 table rows on their own lines, got {len(lines)}'


def test_the_last_paragraph_does_not_run_into_the_table(tmp_path):
    """The += with no separator fused a sentence to the first table header."""
    text = _extract(_doc_with_table(str(tmp_path / 'cv.docx')))
    assert 'Experience follows.Employer' not in text.replace(' ', ''), (
        'the final paragraph is concatenated onto the first table row'
    )
    assert 'Experience follows.' in text


def test_empty_cells_do_not_produce_dangling_separators(tmp_path):
    """A half-filled row is common in real CVs and must not become ' |  | '."""
    p = str(tmp_path / 'sparse.docx')
    d = Document()
    t = d.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = 'Acme Corp'
    t.rows[0].cells[1].text = ''
    t.rows[0].cells[2].text = '2019-2023'
    d.save(p)
    text = _extract(p)
    assert '|  |' not in text, f'empty cell left a dangling separator: {text!r}'
    assert 'Acme Corp | 2019-2023' in text
